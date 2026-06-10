import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

if len(sys.argv) < 3:
    print("Usage: python export_article.py <slug> <platform>")
    print("  platform: wechat | zhihu")
    sys.exit(1)

slug = sys.argv[1]
platform = sys.argv[2].lower()
if platform not in ("wechat", "zhihu"):
    print("platform must be 'wechat' or 'zhihu'")
    sys.exit(1)

ARTICLES = Path(__file__).parent.parent.parent.parent / "ongoing"
# Find latest version directory
versions_dir = ARTICLES / slug / "2-draft"
if not versions_dir.exists():
    print(f"No versions found for {slug}", file=sys.stderr)
    sys.exit(1)
v_dirs = sorted([d for d in versions_dir.iterdir() if d.is_dir() and d.name.startswith("v")],
                key=lambda d: int(d.name[1:]) if d.name[1:].isdigit() else 0)
if not v_dirs:
    print(f"No version directories found for {slug}", file=sys.stderr)
    sys.exit(1)
latest_v = v_dirs[-1]

# Walk back versions to find platform file (inherit rule)
def find_in_versions(filename):
    for vd in reversed(v_dirs):
        f = vd / filename
        if f.exists():
            return f
    return None

MD_FILE = find_in_versions(f"{platform}.md")
if not MD_FILE:
    print(f"Platform file {platform}.md not found in any version", file=sys.stderr)
    sys.exit(1)
IMG_BASE = ARTICLES / slug / "images"

text = MD_FILE.read_text(encoding="utf-8")

_title_match = re.search(r"^#{1,2}\s+(.+)$", text, re.MULTILINE)
_title = _title_match.group(1).strip() if _title_match else slug
_safe_title = re.sub(r'[\\/*?:"<>|]', "_", _title)
OUT_FILE = ARTICLES / slug / f"{_safe_title}.docx"

doc = Document()

style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(11)

HEADING_SIZES = {1: Pt(22), 2: Pt(18), 3: Pt(14), 4: Pt(12)}
HEADING_COLORS = {
    1: RGBColor(0x1a, 0x1a, 0x1a),
    2: RGBColor(0x1a, 0x62, 0xb7),
    3: RGBColor(0x2e, 0x86, 0x48),
    4: RGBColor(0x6b, 0x6b, 0x6b),
}


def set_font(run, bold=False, code=False):
    run.font.name = "微软雅黑"
    run.font.size = Pt(12) if code else Pt(11)
    run.bold = bold
    run.italic = code


def add_heading(doc, text, level):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run.font.size = HEADING_SIZES.get(level, Pt(12))
    run.font.color.rgb = HEADING_COLORS.get(level, RGBColor(0x1a, 0x1a, 0x1a))
    run.bold = True


def add_code_block(doc, code, lang=""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F0F0")
    pPr.append(shd)
    run = p.add_run(code)
    run.font.name = "微软雅黑"
    run.font.size = Pt(12)
    run.italic = True


def add_image(doc, img_path):
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))
    else:
        p = doc.add_paragraph(f"[图片缺失: {img_path.name}]")
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


def set_cell_bg(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def add_table(doc, rows):
    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows) - 1, cols=col_count)
    table.style = "Table Grid"
    hdr = table.rows[0]
    for j, cell_text in enumerate(rows[0]):
        cell = hdr.cells[j]
        cell.text = cell_text.strip()
        set_cell_bg(cell, "D9E1F2")
        for run in cell.paragraphs[0].runs:
            run.font.name = "微软雅黑"
            run.font.size = Pt(10)
            run.bold = True
    for i, row in enumerate(rows[2:], start=1):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = cell_text.strip()
            for run in cell.paragraphs[0].runs:
                run.font.name = "微软雅黑"
                run.font.size = Pt(10)


def parse_table_row(line):
    return line.strip().strip("|").split("|")


def is_table_separator(line):
    return bool(re.match(r"^\|[-| :]+\|$", line.strip()))


def add_inline_paragraph(doc, line):
    p = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))", line)
    for part in parts:
        if part is None:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            set_font(run, code=True)
        elif re.match(r"\[[^\]]+\]\([^)]+\)", part):
            link_text = re.match(r"\[([^\]]+)\]", part).group(1)
            run = p.add_run(link_text)
            set_font(run)
        else:
            run = p.add_run(part)
            set_font(run)


lines = text.splitlines()
i = 0
in_code = False
code_lines = []
code_lang = ""

while i < len(lines):
    line = lines[i]

    img_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line)
    if img_match:
        add_image(doc, IMG_BASE / img_match.group(2).split("/")[-1])
        i += 1
        continue

    if line.startswith("```"):
        if not in_code:
            in_code = True
            code_lang = line[3:].strip()
            code_lines = []
        else:
            in_code = False
            add_code_block(doc, "\n".join(code_lines), code_lang)
        i += 1
        continue

    if in_code:
        code_lines.append(line)
        i += 1
        continue

    h_match = re.match(r"^(#{1,4})\s+(.*)", line)
    if h_match:
        add_heading(doc, h_match.group(2), len(h_match.group(1)))
        i += 1
        continue

    if re.match(r"^---+$", line.strip()):
        doc.add_paragraph("─" * 40)
        i += 1
        continue

    if line.strip() == "":
        i += 1
        continue

    # Markdown table (zhihu only)
    if platform == "zhihu" and line.strip().startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
        table_rows = [parse_table_row(line)]
        i += 1
        table_rows.append(parse_table_row(lines[i]))
        i += 1
        while i < len(lines) and lines[i].strip().startswith("|"):
            table_rows.append(parse_table_row(lines[i]))
            i += 1
        add_table(doc, table_rows)
        continue

    ol_match = re.match(r"^(\d+)\.\s+(.*)", line)
    if ol_match:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(ol_match.group(2))
        set_font(run)
        i += 1
        continue

    ul_match = re.match(r"^[-*]\s+(.*)", line)
    if ul_match:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(ul_match.group(1))
        set_font(run)
        i += 1
        continue

    add_inline_paragraph(doc, line)
    i += 1

doc.save(str(OUT_FILE))
print(f"Saved: {OUT_FILE}")
