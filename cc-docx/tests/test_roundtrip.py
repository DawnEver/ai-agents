import sys
import tempfile
import unittest
from datetime import date
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

import docx2md  # noqa: E402
import md2docx  # noqa: E402


class RoundTripSafetyTests(unittest.TestCase):
    def test_default_output_is_in_project_out_directory(self):
        md_path = ROOT / "workspace" / "ongoing" / "260805-example" / "draft.md"
        template = ROOT / "templates" / "Form.docx"

        expected = md_path.parent / "out" / f"Form-{date.today():%y%m%d}.docx"
        self.assertEqual(md2docx.default_output_path(md_path, template), expected)

    def test_document_round_trips_except_transcript_filename(self):
        with tempfile.TemporaryDirectory() as directory:
            template = Path(directory) / "template.docx"
            first_md = Path(directory) / "first.md"
            output = Path(directory) / "output.docx"
            second_md = Path(directory) / "second.md"
            doc = Document()
            paragraph = doc.add_paragraph("Plain ")
            bold = paragraph.add_run("bold")
            bold.bold = True
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "A"
            table.cell(0, 1).text = "B"
            table.cell(1, 0).text = "1"
            table.cell(1, 1).text = "2"
            doc.save(template)

            first_md.write_text(docx2md.docx2md(str(template)), encoding="utf-8")
            md2docx.md2docx(str(first_md), str(template), str(output))
            second_md.write_text(docx2md.docx2md(str(output)), encoding="utf-8")

            first_lines = first_md.read_text(encoding="utf-8").splitlines()[1:]
            second_lines = second_md.read_text(encoding="utf-8").splitlines()[1:]
            self.assertEqual(first_lines, second_lines)

    def test_table_separator_has_the_effective_column_count(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"

        self.assertEqual(docx2md._table_md(table), ["| A | B |", "|---|---|"])

    def test_render_refuses_to_overwrite_template(self):
        with tempfile.TemporaryDirectory() as directory:
            template = Path(directory) / "template.docx"
            transcript = Path(directory) / "template.md"
            doc = Document()
            doc.add_paragraph("Original")
            doc.save(template)
            docx2md.docx2md(str(template))
            transcript.write_text("<!-- ccx1 --> Changed\n", encoding="utf-8")

            with self.assertRaisesRegex(ValueError, "template"):
                md2docx.md2docx(str(transcript), str(template), str(template))

    def test_render_rejects_unknown_anchor(self):
        with tempfile.TemporaryDirectory() as directory:
            template = Path(directory) / "template.docx"
            transcript = Path(directory) / "template.md"
            output = Path(directory) / "output.docx"
            doc = Document()
            doc.add_paragraph("Original")
            doc.save(template)
            docx2md.docx2md(str(template))
            transcript.write_text("<!-- ccx999 --> Changed\n", encoding="utf-8")

            with self.assertRaisesRegex(ValueError, "ccx999"):
                md2docx.md2docx(str(transcript), str(template), str(output))
            self.assertFalse(output.exists())

    def test_render_rejects_anchor_kind_mismatch(self):
        with tempfile.TemporaryDirectory() as directory:
            template = Path(directory) / "template.docx"
            transcript = Path(directory) / "template.md"
            output = Path(directory) / "output.docx"
            doc = Document()
            doc.add_paragraph("Original")
            doc.save(template)
            docx2md.docx2md(str(template))
            transcript.write_text(
                "<!-- ccx1 -->\n| A |\n|---|\n| B |\n", encoding="utf-8"
            )

            with self.assertRaisesRegex(ValueError, "ccx1"):
                md2docx.md2docx(str(transcript), str(template), str(output))
            self.assertFalse(output.exists())


if __name__ == "__main__":
    unittest.main()
