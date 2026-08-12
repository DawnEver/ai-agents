#!/usr/bin/env python3
"""Generic structured-content → Word docx builder.

Loads a content module (Python file) and renders it as a clean Word document
with a 4-level nested-bullet hierarchy (distinct glyphs per level), and — when a
baseline document is given — a colour-coded review of changes against that baseline.

Content module format (see workspace/260812-JoR/jor_content.py for an example):
    TITLE   = "Document title"
    CONTENT = [ (kind, text, chg, change_desc), ... ]
        kind : l1|l2|l2p|l3|l4|p|table
        chg  : kept|edited|new|moved     (used only for baseline-relative review)

Usage:
    python build_docx.py <content.py> <out_dir> [--baseline <baseline.docx>] [--name <stem>]
Outputs (in <out_dir>):
    <name>.docx            clean document
    <name>_review.docx     colour-coded review vs <baseline>  (only with --baseline)

The baseline-relative review classifies each content item against the baseline
document by fuzzy text matching: high similarity → KEPT, moderate → EDITED,
none → NEW. Items with low similarity to everything are marked NEW.
"""
import argparse
import datetime
import difflib
import importlib.util
import os
import re
import sys

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- tunable presentation -------------------------------------------------
BULLET_LEVEL = {"l1": 0, "l2": 1, "l2p": 1, "l3": 2, "l4": 3}
GLYPHS = ["•", "–", "▪", "·"]   # • – ▪ ·
INDENTS = [540, 900, 1260, 1620]                    # twips per level
FILL = {"new": "E2EFDA", "edited": "FFF2CC", "moved": "DDEBF7", "kept": None}
LABEL = {"new": "NEW", "edited": "EDITED", "moved": "MOVED", "kept": "KEPT"}
KEPT_HEADINGS = ("Personnel",)   # headings never flagged as changed


# --------------------------------------------------------------------------
# numbering (4-level bullets with distinct glyphs)
# --------------------------------------------------------------------------
def setup_numbering(doc):
    num_part = doc.part.numbering_part.element
    abstractNum = OxmlElement("w:abstractNum")
    abstractNum.set(qn("w:abstractNumId"), "100")
    for level in range(4):
        lvl = OxmlElement("w:lvl"); lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start"); start.set(qn("w:val"), "1")
        numFmt = OxmlElement("w:numFmt"); numFmt.set(qn("w:val"), "bullet")
        lvlText = OxmlElement("w:lvlText"); lvlText.set(qn("w:val"), GLYPHS[level])
        rPr = OxmlElement("w:rPr"); rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Calibri"); rFonts.set(qn("w:hAnsi"), "Calibri")
        rPr.append(rFonts)
        pPr = OxmlElement("w:pPr"); ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(INDENTS[level])); ind.set(qn("w:hanging"), "360")
        pPr.append(ind)
        for el in (start, numFmt, lvlText, rPr, pPr):
            lvl.append(el)
        abstractNum.append(lvl)
    num_part.append(abstractNum)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), "100")
    abId = OxmlElement("w:abstractNumId"); abId.set(qn("w:val"), "100")
    num.append(abId)
    num_part.append(num)


def new_document():
    doc = Document()
    n = doc.styles["Normal"]
    n.font.name = "Calibri"; n.font.size = Pt(11); n.paragraph_format.space_after = Pt(4)
    setup_numbering(doc)
    return doc


def shade(p, fill):
    if not fill:
        return
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_par(doc, text, bold=False, centered=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.bold = bold
    if size:
        r.font.size = size
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_bullet(doc, text, level, bold=False):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), str(level))
    numId = OxmlElement("w:numId"); numId.set(qn("w:val"), "100")
    numPr.append(ilvl); numPr.append(numId)
    pPr.insert(0, numPr)
    r = p.add_run(text); r.font.bold = bold
    return p


# --------------------------------------------------------------------------
# content rendering
# --------------------------------------------------------------------------
def render_item(doc, kind, val):
    """Add one content item to `doc`. Returns the added paragraph (or None)."""
    if kind == "table":
        rows = val
        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = "Table Grid"
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                c = t.cell(i, j); c.text = ""
                p = c.paragraphs[0]; r = p.add_run(cell); r.font.size = Pt(10)
                if i == 0:
                    r.font.bold = True
        return None
    if kind == "p":
        return add_par(doc, val)
    return add_bullet(doc, val, BULLET_LEVEL[kind], bold=(kind in ("l1", "l2")))


def build_clean(title, content):
    doc = new_document()
    add_par(doc, title, bold=True, centered=True, size=Pt(16))
    for item in content:
        kind, val = item[0], item[1]
        render_item(doc, kind, val)
    return doc


# --------------------------------------------------------------------------
# baseline-relative review
# --------------------------------------------------------------------------
def load_baseline_texts(path):
    from docx import Document as _Doc
    W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    d = _Doc(path)
    return [''.join(x.text or '' for x in p._p.iter(W + 't')).strip()
            for p in d.paragraphs if ''.join(x.text or '' for x in p._p.iter(W + 't')).strip()]


def norm(t):
    return re.sub(r'[^a-z0-9]+', ' ', t.lower())


def classify_vs_baseline(kind, val, baseline_texts):
    """Return (chg, description) for one content item against a baseline."""
    if kind in ("l1", "l2", "l2p") and val in KEPT_HEADINGS:
        return "kept", "kept from baseline"
    nv = norm(val)
    best = 0.0
    for t in baseline_texts:
        best = max(best, difflib.SequenceMatcher(None, nv, norm(t)).ratio())
    if best > 0.82:
        return "kept", "kept from baseline (wording ~ unchanged)"
    if best > 0.45:
        return "edited", "rewritten / updated vs baseline"
    return "new", "new vs baseline"


def build_review(title, content, baseline_path, baseline_label="baseline"):
    doc = new_document()
    baseline_texts = load_baseline_texts(baseline_path)

    add_par(doc, f"REVIEW — changes vs {baseline_label}", bold=True, size=Pt(14))
    add_par(doc, "Body colour-coded by change type:")
    for lab, fill, meaning in [("NEW", FILL["new"], "not present in baseline"),
                               ("EDITED", FILL["edited"], "rewritten / updated from baseline"),
                               ("KEPT", None, "unchanged from baseline")]:
        p = doc.add_paragraph(); r = p.add_run(f"■ {lab}  = {meaning}")
        if fill:
            shade(p, fill)
    add_par(doc, "")
    add_par(doc, "Change log:", bold=True)

    log_rows = [["#", "Type", "Item", "What changed"]]
    n = 0
    for kind, val, _c0, _d0 in content:
        c, d = classify_vs_baseline(kind, val, baseline_texts)
        if c == "kept":
            continue
        n += 1
        if kind == "l1":
            label = val
        elif kind == "l4":
            label = val[:42] + ("..." if len(val) > 42 else "")
        else:
            label = (val.split("[")[0].strip()[:42] if "[" in val else val[:42])
        log_rows.append([str(n), LABEL[c], label, d])
    t = doc.add_table(rows=len(log_rows), cols=4); t.style = "Table Grid"
    for i, row in enumerate(log_rows):
        for j, cell in enumerate(row):
            c = t.cell(i, j); c.text = ""
            p = c.paragraphs[0]; r = p.add_run(cell); r.font.size = Pt(8.5)
            if i == 0:
                r.font.bold = True
            elif i > 0 and row[1] in ("NEW", "EDITED", "MOVED"):
                shade_cell(c, FILL[row[1].lower()])
    doc.add_page_break()

    add_par(doc, title, bold=True, centered=True, size=Pt(16))
    for kind, val, _c0, _d0 in content:
        if kind == "table":
            render_item(doc, kind, val)
            continue
        c, d = classify_vs_baseline(kind, val, baseline_texts)
        p = render_item(doc, kind, val)
        if p is not None:
            shade(p, FILL[c])
    return doc


# --------------------------------------------------------------------------
def word_count(title, content):
    joined = [title] + [val for kind, val, *_ in content if kind != "table"] \
             + [" ".join(r) for kind, val, *_ in content if kind == "table" for r in val]
    return len(re.findall(r"\b\w+\b", " ".join(joined)))


def load_content_module(path):
    spec = importlib.util.spec_from_file_location("usercontent", os.path.abspath(path))
    mod = importlib.util.module_from_spec(spec)
    sys.modules["usercontent"] = mod
    spec.loader.exec_module(mod)
    return mod.TITLE, mod.CONTENT


def main():
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("content", help="content module .py path")
    ap.add_argument("out_dir", help="output directory")
    ap.add_argument("--baseline", help="baseline .docx for a colour-coded review")
    ap.add_argument("--name", help="output stem (default: derived from content filename)")
    ap.add_argument("--label", default="baseline", help="label shown in the review header")
    args = ap.parse_args()

    title, content = load_content_module(args.content)
    name = args.name or os.path.splitext(os.path.basename(args.content))[0]
    os.makedirs(args.out_dir, exist_ok=True)

    build_clean(title, content).save(os.path.join(args.out_dir, f"{name}.docx"))
    print(f"clean:  {os.path.join(args.out_dir, name + '.docx')}")
    if args.baseline:
        build_review(title, content, args.baseline, args.label).save(
            os.path.join(args.out_dir, f"{name}_review.docx"))
        print(f"review: {os.path.join(args.out_dir, name + '_review.docx')}")
    print(f"word count (incl. table cells): {word_count(title, content)}")


if __name__ == "__main__":
    main()
