# -*- coding: utf-8 -*-
"""docx → markdown transcript with patch-back anchors.

Transcribes every body block (paragraphs, lists, tables) of a .docx into a
markdown working copy, and stamps an invisible bookmark pair (`ccxN`) around
each block in the docx so `md2docx.py` can later patch content back into the
original template without disturbing its layout/styles.

Usage:  python scripts/docx2md.py <input.docx> [output.md]
"""
import io
import os
import re
import sys
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

ANCHOR_PREFIX = "ccx"
# heading style names python-docx exposes as 'Heading 1'..'Heading 9'
HEADING_RE = re.compile(r"^Heading\s*([1-9])$", re.I)


# ---------------------------------------------------------------- xml helpers

def _ns(tag):
    return qn(tag)


def _iter_runs_in_order(p_elm):
    """Yield w:r runs and w:hyperlink blocks in true document order.
    w:ins (track-changes insertions) are transparent — their runs are
    yielded in place so revisions read back as plain text."""
    for child in p_elm.iterchildren():
        if child.tag in (_ns("w:r"), _ns("w:hyperlink")):
            yield child
        elif child.tag == _ns("w:ins"):
            for r in child.iterchildren():
                if r.tag in (_ns("w:r"), _ns("w:hyperlink")):
                    yield r


def _run_text(r_elm):
    return "".join(t.text or "" for t in r_elm.findall(_ns("w:t")))


def _run_attrs(r_elm):
    rpr = r_elm.find(_ns("w:rPr"))
    attrs = {"b": False, "i": False, "u": False, "mono": False, "hl": False}
    if rpr is None:
        return attrs
    if rpr.find(_ns("w:b")) is not None:
        attrs["b"] = True
    if rpr.find(_ns("w:i")) is not None:
        attrs["i"] = True
    if rpr.find(_ns("w:u")) is not None:
        attrs["u"] = True
    # yellow highlight → `==text==` (round-trip pair for md2docx)
    hl = rpr.find(_ns("w:highlight"))
    if hl is not None and (hl.get(_ns("w:val")) or "yellow") == "yellow":
        attrs["hl"] = True
    # Courier-like / code faces → backticks
    fonts = [f.get(_ns("w:ascii")) or "" for f in rpr.findall(_ns("w:rFonts"))]
    if any("Courier" in f or "Consolas" in f or "Monaco" in f for f in fonts):
        attrs["mono"] = True
    return attrs


def _escape_md(text):
    # escape characters that would be misread as markdown markup
    return text.replace("\\", "\\\\").replace("|", "\\|").replace("*", "\\*")


def _inline_md(elm, part):
    """→ [(fmt, text)] for a run-or-hyperlink element. fmt in '', 'b', 'i',
    'code', or ('link', url)."""
    if elm.tag == _ns("w:hyperlink"):
        rid = elm.get(_ns("r:id"))
        url = ""
        if rid and part is not None:
            rel = part.rels[rid]
            url = rel.target_ref if rel is not None else ""
        tokens = []
        for r in elm.findall(_ns("w:r")):
            tokens.extend(_inline_md(r, part))
        if url:
            return [(("link", url), "".join(t[1] for t in tokens))]
        return tokens
    text = _run_text(elm)
    if not text:
        return []
    attrs = _run_attrs(elm)
    if attrs["hl"]:
        fmt = "hl"
        if attrs["b"]:
            fmt = "hl+b"
        elif attrs["i"]:
            fmt = "hl+i"
        elif attrs["mono"]:
            fmt = "hl+code"
    elif attrs["b"]:
        fmt = "b"
    elif attrs["i"]:
        fmt = "i"
    elif attrs["mono"]:
        fmt = "code"
    else:
        fmt = ""
    return [(fmt, text)]


def _para_text_md(p):
    """Full inline markdown for a paragraph, including hyperlinks.
    Adjacent runs with identical formatting are merged first (Word splits
    bold text into per-space runs, which would otherwise produce `****`)."""
    tokens = []
    for elm in _iter_runs_in_order(p._p):
        tokens.extend(_inline_md(elm, p.part))
    merged = []
    for tok in tokens:
        key = tok[0]
        if merged and merged[-1][0] == key:
            merged[-1] = (key, merged[-1][1] + tok[1])
        else:
            merged.append((key, tok[1]))
    out = []
    for fmt, text in merged:
        t = _escape_md(text)
        if fmt == "b":
            t = f"**{t}**"
        elif fmt == "i":
            t = f"*{t}*"
        elif fmt == "code":
            t = f"`{t}`"
        elif fmt == "hl":
            t = f"=={t}=="
        elif fmt == "hl+b":
            t = f"==**{t}**=="
        elif fmt == "hl+i":
            t = f"==*{t}*=="
        elif fmt == "hl+code":
            t = f"==`{t}`=="
        elif isinstance(fmt, tuple) and fmt[0] == "link":
            t = f"[{t}]({fmt[1]})"
        out.append(t)
    return "".join(out)


def _list_info(p):
    """Return (numid, ilvl) if paragraph is in a numbered/bulleted list."""
    ppr = p._p.find(_ns("w:pPr"))
    if ppr is None:
        return None
    numpr = ppr.find(_ns("w:numPr"))
    if numpr is None:
        return None
    ilvl = numpr.find(_ns("w:ilvl"))
    numid = numpr.find(_ns("w:numId"))
    return (
        numid.get(_ns("w:val")) if numid is not None else None,
        int(ilvl.get(_ns("w:val")) or 0) if ilvl is not None else 0,
    )


def _num_format(doc, numid):
    """Is list numbering decimal (ordered) or bullet? Walks numbering part."""
    try:
        numbering = doc.part.numbering_part.numbering_definitions
        num_el = None
        for n in numbering.findall(qn("w:num")):
            if n.get(qn("w:numId")) == numid:
                num_el = n
                break
        if num_el is None:
            return "bullet"
        abs_id = num_el.find(qn("w:abstractNumId")).get(qn("w:val"))
        for an in numbering.findall(qn("w:abstractNum")):
            if an.get(qn("w:abstractNumId")) == abs_id:
                lvl = an.find(qn("w:lvl"))
                if lvl is not None:
                    fmt = lvl.find(qn("w:numFmt"))
                    if fmt is not None and fmt.get(qn("w:val")) == "decimal":
                        return "decimal"
                break
    except Exception:
        pass
    return "bullet"


# ------------------------------------------------------------- table handling

def _table_matrix(table):
    """Grid-indexed matrix: row.cells is indexed by grid column, so merged
    cells (gridSpan / vMerge) appear repeated at every spanned position —
    md ↔ docx map 1:1 positionally with no dedup."""
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            texts = [_para_text_md(p) for p in cell.paragraphs]
            text = " ".join(t for t in texts if t.strip()).strip()
            cells.append(text)
        rows.append(cells)
    return rows


def _table_md(table):
    rows = _table_matrix(table)
    if not rows:
        return []
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    out = []
    out.append("| " + " | ".join(rows[0]) + " |")
    out.append("|" + "---|" * (width + 1))
    for r in rows[1:]:
        out.append("| " + " | ".join(r) + " |")
    return out


# --------------------------------------------------------------- anchoring

def _bookmark_names(root):
    """ALL existing ccx bookmark names under root (paragraph-internal and
    body-level) — collected up front so new ids can never collide with
    bookmarks that appear later in document order."""
    found = set()
    for bm in root.iter(_ns("w:bookmarkStart")):
        n = bm.get(_ns("w:name") or "")
        if n.startswith(ANCHOR_PREFIX):
            found.add(n)
    return found


def _anchor_block(body_elm, block_elm, used_ids):
    """Wrap a body-level element in a ccxN bookmark pair; return the id.
    Reuses the bookmark if the block is already preceded by one."""
    prev = block_elm.getprevious()
    if prev is not None and prev.tag == _ns("w:bookmarkStart"):
        name = prev.get(_ns("w:name") or "")
        if name.startswith(ANCHOR_PREFIX):
            used_ids.add(name)
            return name
    n = 1
    while f"{ANCHOR_PREFIX}{n}" in used_ids:
        n += 1
    aid = f"{ANCHOR_PREFIX}{n}"
    used_ids.add(aid)
    start = body_elm.makeelement(_ns("w:bookmarkStart"), {
        _ns("w:id"): str(n), _ns("w:name"): aid})
    end = body_elm.makeelement(_ns("w:bookmarkEnd"), {_ns("w:id"): str(n)})
    block_elm.addprevious(start)
    block_elm.addnext(end)
    return aid


def _paragraph_anchor(body_elm, p_elm, used_ids):
    """Anchor inside the paragraph so empty paragraphs can be found too.
    Bookmark goes AFTER w:pPr (Word schema order: pPr, then markup/runs).
    Reuses an existing ccx bookmark if the paragraph already has one."""
    for bm in p_elm.findall(_ns("w:bookmarkStart")):
        name = bm.get(_ns("w:name") or "")
        if name.startswith(ANCHOR_PREFIX):
            used_ids.add(name)
            return name
    n = 1
    while f"{ANCHOR_PREFIX}{n}" in used_ids:
        n += 1
    aid = f"{ANCHOR_PREFIX}{n}"
    used_ids.add(aid)
    start = body_elm.makeelement(_ns("w:bookmarkStart"), {
        _ns("w:id"): str(n), _ns("w:name"): aid})
    end = body_elm.makeelement(_ns("w:bookmarkEnd"), {_ns("w:id"): str(n)})
    ppr = p_elm.find(_ns("w:pPr"))
    if ppr is not None:
        ppr.addnext(start)
    else:
        p_elm.insert(0, start)
    p_elm.append(end)
    return aid


# ------------------------------------------------------------- main pipeline

def _header_footer_blocks(doc):
    """Transcribe section headers/footers as blockquote lines, anchored."""
    blocks = []
    for section in doc.sections:
        for label, container in (("HEADER", section.header),
                                 ("FOOTER", section.footer)):
            for p in container.paragraphs:
                if not _para_text_md(p).strip():
                    continue
                blocks.append((label, p))
    return blocks


def docx2md(docx_path):
    doc = Document(docx_path)
    body_elm = doc.element.body
    used_ids = _bookmark_names(body_elm)
    anchors_added = False
    out = []

    out.append(f"<!-- docx2md transcript of: {os.path.basename(docx_path)} -->")
    out.append(f"<!-- anchors: {ANCHOR_PREFIX}N | do not delete anchor comments -->")
    out.append("")

    # headers/footers first (transcript only; rendered back on patch)
    for label, p in _header_footer_blocks(doc):
        aid = _paragraph_anchor(body_elm, p._p, used_ids)
        anchors_added = True
        out.append(f"> **{label}** <!-- {aid} --> {_para_text_md(p)}")
    if _header_footer_blocks(doc):
        out.append("")

    for child in body_elm.iterchildren():
        if child.tag == _ns("w:p"):
            p = Paragraph(child, doc)
            # use the inline-md rendering (ins-aware) for the emptiness test —
            # p.text misses runs inside w:ins (track-changes) elements
            text = _para_text_md(p).strip()
            aid = _paragraph_anchor(body_elm, child, used_ids)
            anchors_added = True
            if not text:
                out.append(f"<!-- {aid} -->")
                out.append("")
                continue
            ppr = child.find(_ns("w:pPr"))
            style_name = ""
            if ppr is not None:
                pstyle = ppr.find(_ns("w:pStyle"))
                if pstyle is not None:
                    style_name = pstyle.get(_ns("w:val")) or ""
            m = HEADING_RE.match(style_name)
            if m:
                lvl = int(m.group(1))
                out.append(f"{'#' * lvl} <!-- {aid} --> {_para_text_md(p)}")
            else:
                li = _list_info(p)
                if li is not None:
                    numid, ilvl = li
                    kind = "1." if _num_format(doc, numid) == "decimal" else "-"
                    indent = "  " * ilvl
                    out.append(f"{indent}{kind} <!-- {aid} --> {_para_text_md(p)}")
                else:
                    out.append(f"<!-- {aid} --> {_para_text_md(p)}")
            out.append("")
        elif child.tag == _ns("w:tbl"):
            table = Table(child, doc)
            aid = _anchor_block(body_elm, child, used_ids)
            anchors_added = True
            out.append(f"<!-- {aid} -->")
            out.extend(_table_md(table))
            out.append("")
    if anchors_added:
        doc.save(docx_path)  # persist the ccxN bookmarks into the docx
    return "\n".join(out).rstrip() + "\n"


def main():
    if len(sys.argv) < 2:
        sys.exit(__doc__)
    src = sys.argv[1]
    if len(sys.argv) > 2:
        dst = sys.argv[2]
    else:
        dst = os.path.splitext(os.path.basename(src))[0] + ".md"
    text = docx2md(src)
    os.makedirs(os.path.dirname(os.path.abspath(dst)) or ".", exist_ok=True)
    with io.open(dst, "w", encoding="utf-8") as f:
        f.write(text)
    print(f"transcribed: {src} -> {dst}")


if __name__ == "__main__":
    main()
