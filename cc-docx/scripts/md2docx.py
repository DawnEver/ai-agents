# -*- coding: utf-8 -*-
"""markdown → docx patch-back renderer.

Reads a `docx2md.py` transcript (.md), walks the ORIGINAL template .docx
(which still carries the `ccxN` bookmarks), and patches each anchored block
with the md's content. Original layout and styles survive; unanchored
(new) blocks are inserted after the previous anchored block.

Usage:  python scripts/md2docx.py <input.md> <template.docx> [output.docx] [--track-changes]

`--track-changes` (review mode): rewritten blocks are wrapped in w:ins
revision elements (author "AI Agent") so the user can accept or
reject AI-added content in Word's Review pane. Blocks whose md content
equals the template's current text are left untouched (no revision).
"""
import copy
import datetime
import io
import os
import re
import sys
from pathlib import Path
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

ANCHOR_PREFIX = "ccx"
ANCHOR_ONLY = re.compile(r"^<!--\s*(ccx\d+)\s*-->$")
HEADING_LINE = re.compile(r"^(#{1,6})\s*(?:<!--\s*(ccx\d+)\s*-->)?\s*(.*)$")
LIST_LINE = re.compile(r"^\s*(?:[-*]|\d+\.)\s+(?:<!--\s*(ccx\d+)\s*-->)?\s*(.*)$")
HEADER_LINE = re.compile(
    r"^>\s*\*\*(HEADER|FOOTER)\*\*\s*(?:<!--\s*(ccx\d+)\s*-->)?\s*(.*)$")

# Track-changes (review mode): inserted content is wrapped in w:ins so the
# user can accept/reject it in Word's Review pane. Enabled with --track-changes.
TRACK_FLAG = "--track-changes"
AUTHOR = "AI Agent"
DATE = "2026-08-05T00:00:00Z"

# ------------------------------------------------------------ inline markdown

TOKEN_RE = re.compile(
    r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]*\]\([^)]*\)|[^*`]+)")
# `==text==` yellow highlight — split BEFORE the plain tokeniser, whose
# greedy catch-all would otherwise swallow it mid-text.
HL_RE = re.compile(r"==(.*?)==")


def parse_inline(text):
    """→ [(text, fmt)] with fmt in '', 'b', 'i', 'code', 'hl', 'hl+b',
    'hl+i', 'hl+code', or ('link', url). `==text==` renders as yellow
    highlight in Word; `==**text**==` combines highlight + bold."""
    tokens = []
    pos = 0
    for m in HL_RE.finditer(text):
        if m.start() > pos:
            _parse_inline_plain(text[pos:m.start()], tokens)
        inner = m.group(1)
        sub = parse_inline(inner)
        if len(sub) == 1 and sub[0][1] in ("b", "i", "code"):
            tokens.append((sub[0][0], "hl+" + sub[0][1]))
        else:
            tokens.append((inner, "hl"))
        pos = m.end()
    if pos < len(text):
        _parse_inline_plain(text[pos:], tokens)
    return tokens


def _parse_inline_plain(text, tokens):
    for m in TOKEN_RE.finditer(text):
        tok = m.group(0)
        if tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
            tokens.append((tok[2:-2], "b"))
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            tokens.append((tok[1:-1], "i"))
        elif tok.startswith("`") and tok.endswith("`"):
            tokens.append((tok[1:-1], "code"))
        else:
            link = re.match(r"\[([^\]]*)\]\(([^)]*)\)", tok)
            if link and link.group(0) == tok:
                tokens.append((link.group(1), ("link", link.group(2))))
            else:
                tokens.append((tok, ""))


def _make_element(elm, tag, attrs=None):
    e = elm.makeelement(qn(tag), attrs or {})
    return e


def _new_run(paragraph, text, fmt, link_part, ins_ctx):
    """Append one run. `ins_ctx` (dict with 'n') wraps the run in a w:ins
    revision element when track-changes is on. w:t keeps leading/trailing
    spaces with xml:space="preserve" — without it Word silently drops
    run-boundary spaces (e.g. 'adding ==GBP 200,000== by' → 'addingGBP…by')."""
    elm = paragraph._p
    holder = elm
    if ins_ctx is not None:
        ins_ctx["n"] += 1
        ins = _make_element(elm, "w:ins", {
            qn("w:id"): str(ins_ctx["n"]),
            qn("w:author"): AUTHOR,
            qn("w:date"): DATE})
        elm.append(ins)
        holder = ins
    if isinstance(fmt, tuple) and fmt[0] == "link":
        url = fmt[1]
        rel = link_part.relate_to(url, RT.HYPERLINK, is_external=True)
        link = _make_element(elm, "w:hyperlink", {qn("r:id"): rel.rId})
        run = _make_element(elm, "w:r")
        link.append(run)
        holder.append(link)
    else:
        run = _make_element(elm, "w:r")
        holder.append(run)
    hl = fmt in ("hl", "hl+b", "hl+i", "hl+code")
    sub = fmt[3:] if hl else fmt
    if hl or sub in ("b", "i", "code"):
        rpr = _make_element(elm, "w:rPr")
        if hl:
            rpr.append(_make_element(elm, "w:highlight", {qn("w:val"): "yellow"}))
        if sub == "b":
            rpr.append(_make_element(elm, "w:b"))
        elif sub == "i":
            rpr.append(_make_element(elm, "w:i"))
        elif sub == "code":
            rpr.append(_make_element(elm, "w:rFonts", {
                qn("w:ascii"): "Consolas", qn("w:hAnsi"): "Consolas"}))
        run.append(rpr)
    t = _make_element(elm, "w:t")
    t.text = text
    if text != text.strip():
        t.set(qn("xml:space"), "preserve")
    run.append(t)


def _plain_text(elm):
    """Concatenated text of all w:t under elm (template current content)."""
    return "".join(t.text or "" for t in elm.iter(qn("w:t")))


def _md_plain(text):
    """Strip inline markdown markers from an md block's text so it can be
    compared with the template's plain text. Without this, every block whose
    md carries `**bold**`/`==highlight==`/link markers looks "changed" and
    gets rewritten (and revision-marked) even though the text is identical."""
    return "".join(t for t, _ in parse_inline(text))


def _norm_ws(s):
    """Normalise whitespace and curly quotes so unchanged template text
    isn't misjudged as changed. The docx carries double spaces, NBSP and
    curly quotes (' ' " ") that the md copy writes as single spaces and
    straight quotes; without this, untouched template lines (e.g. "time.
    We have already added..." or "Associate's holidays") get revision-
    marked even though nothing changed."""
    s = re.sub(r"\s+", " ", s).strip()
    return (s.replace("’", "'").replace("‘", "'")
             .replace("“", '"').replace("”", '"'))


def _has_bold(elm):
    return any(r.find(qn("w:rPr")) is not None and
               r.find(qn("w:rPr")).find(qn("w:b")) is not None
               for r in elm.iter(qn("w:r")))


def _unchanged(template_elm, md_text):
    """Same plain text (whitespace/quote-normalised) AND no formatting the
    md adds that the template text doesn't carry. The second clause keeps
    answer marks like "Select one: Yes / **No**" — same text, but the bold
    is an intentional change — while template headings (already bold in the
    docx) and untouched instruction lines stay pristine."""
    if _norm_ws(_plain_text(template_elm)) != _norm_ws(_md_plain(md_text)):
        return False
    if "**" in md_text and not _has_bold(template_elm):
        return False
    return True


def set_paragraph_md(p, text, link_part, track=False):
    """Replace a paragraph's runs with parsed inline markdown.
    Keeps w:pPr and any ccx bookmark pair (the round-trip anchors).
    With track-changes on, blocks whose content is unchanged are left
    untouched (no revision); changed blocks are rewritten as w:ins."""
    elm = p._p
    if track and _unchanged(elm, text):
        return  # unchanged — template text stays pristine, no revision
    ins_ctx = {"n": 0} if track else None
    for child in list(elm):
        if child.tag in (qn("w:pPr"), qn("w:bookmarkStart"),
                         qn("w:bookmarkEnd")):
            continue
        elm.remove(child)
    for text_i, fmt in parse_inline(text):
        _new_run(p, text_i, fmt, link_part, ins_ctx)


def _cell_plain(cell):
    """All paragraphs of a cell joined — a template cell can hold one
    sentence split across several paragraphs (e.g. the T&D description),
    so comparing only paras[0] would misjudge it as changed."""
    return " ".join(_plain_text(p._p) for p in cell.paragraphs)


def set_cell_text(cell, text, link_part, track=False):
    paras = cell.paragraphs
    if track and _norm_ws(_cell_plain(cell)) == _norm_ws(_md_plain(text)):
        return  # unchanged cell — leave as-is (multi-paragraph layout kept)
    for extra in paras[1:]:
        extra._p.getparent().remove(extra._p)
    set_paragraph_md(paras[0], text, link_part, track=track)


# -------------------------------------------------------------- md parsing

def _parse_table(lines):
    rows = []
    for line in lines:
        cells = [c.strip().replace(r"\|", "|") for c in line.strip().strip("|").split("|")]
        # separator row = every cell is dashes; empty cells are DATA, not separators
        if cells and all(c and re.fullmatch(r":?-{3,}:?", c) for c in cells):
            continue
        rows.append(cells)
    return rows


PARA_ANCHOR = re.compile(r"^<!--\s*(ccx\d+)\s*-->\s*(.*)$")
PLAIN_COMMENT = re.compile(r"^<!--.*-->$")


def _parse_md_full(text):
    """→ list of block dicts:
    {kind: 'empty'|'para'|'heading'|'list'|'header'|'table',
     anchor: id|None, text?, level?, rows?}
    Tables are consumed contiguously (index-based loop)."""
    lines = text.splitlines()
    blocks = []
    pending = None
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        m = ANCHOR_ONLY.match(line)
        if m:
            pending = m.group(1)
            i += 1
            continue
        if not line.strip():
            if pending:
                blocks.append({"kind": "empty", "anchor": pending})
                pending = None
            i += 1
            continue
        if line.lstrip().startswith("|"):
            tbl = [line]
            i += 1
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                tbl.append(lines[i].rstrip())
                i += 1
            blocks.append({"kind": "table", "anchor": pending,
                           "rows": _parse_table(tbl)})
            pending = None
            continue
        m = HEADER_LINE.match(line)
        if m:
            blocks.append({"kind": "header", "anchor": m.group(2),
                           "text": m.group(3)})
            i += 1
            continue
        m = HEADING_LINE.match(line)
        if m:
            blocks.append({"kind": "heading", "anchor": m.group(2),
                           "level": len(m.group(1)), "text": m.group(3)})
            i += 1
            continue
        m = LIST_LINE.match(line)
        if m:
            blocks.append({"kind": "list", "anchor": m.group(1),
                           "text": m.group(2)})
            i += 1
            continue
        m = PARA_ANCHOR.match(line)  # extraction format: <!-- ccxN --> text
        if m:
            blocks.append({"kind": "para", "anchor": m.group(1),
                           "text": m.group(2)})
            i += 1
            continue
        if PLAIN_COMMENT.match(line):
            i += 1  # metadata comment (e.g. transcript header) — skip
            continue
        blocks.append({"kind": "para", "anchor": pending, "text": line})
        pending = None
        i += 1
    if pending:
        blocks.append({"kind": "empty", "anchor": pending})
    return blocks


# ---------------------------------------------------------------- anchors

def _block_anchors(root, anchors):
    """Find ccxN bookmarkStart elements under `root`; record target block."""
    for bm in root.iter(qn("w:bookmarkStart")):
        name = bm.get(qn("w:name") or "")
        if not name.startswith(ANCHOR_PREFIX):
            continue
        parent = bm.getparent()
        if parent.tag == qn("w:p"):
            anchors[name] = ("p", parent)
        else:
            nxt = bm.getnext()
            if nxt is not None and nxt.tag in (qn("w:p"), qn("w:tbl")):
                anchors[name] = (nxt.tag.rsplit("}", 1)[-1], nxt)


def build_anchor_map(doc):
    anchors = {}
    _block_anchors(doc.element.body, anchors)
    for section in doc.sections:
        for container in (section.header, section.footer):
            _block_anchors(container._element, anchors)
    return anchors


# ------------------------------------------------------------- table sync

def sync_table(table, md_rows, link_part, track=False):
    """Fill table from md rows. Grid-indexed: row.cells maps 1:1 to md cells
    (merged cells are written at every spanned position — same tc, same text)."""
    nrows = len(md_rows)
    while len(table.rows) < nrows:
        last = table.rows[-1]._tr
        last.addnext(copy.deepcopy(last))
    while len(table.rows) > nrows:
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)
    for r in range(nrows):
        cells = table.rows[r].cells
        for c in range(min(len(cells), len(md_rows[r]))):
            set_cell_text(cells[c], md_rows[r][c], link_part, track=track)


# --------------------------------------------------------------- insert

def insert_block(doc, body_elm, cursor, block, link_part, track=False):
    """Insert an unanchored block after cursor; returns the new element.
    With track-changes on, the whole new block is a w:ins revision."""
    if block["kind"] == "table":
        t = doc.add_table(rows=len(block["rows"]),
                          cols=max((len(r) for r in block["rows"]), default=1))
        if "Table Grid" in [s.name for s in doc.styles]:
            t.style = doc.styles["Table Grid"]
        tbl = t._tbl
        body_elm.remove(tbl)
        cursor.addnext(tbl)
        sync_table(t, block["rows"], link_part, track=track)
        return tbl
    # paragraph-like
    if block["kind"] == "heading":
        style = f"Heading {block.get('level', 1)}"
    else:
        style = None
    p = doc.add_paragraph("", style=style) if style else doc.add_paragraph()
    body_elm.remove(p._p)
    cursor.addnext(p._p)
    ins_ctx = {"n": 0} if track else None
    for text_i, fmt in parse_inline(block.get("text", "")):
        _new_run(p, text_i, fmt, link_part, ins_ctx)
    return p._p


# ----------------------------------------------------------------- main

def default_output_path(md_path, docx_path, today=None):
    """Return a traceable output path inside the transcript's project."""
    stem = os.path.splitext(os.path.basename(os.fspath(docx_path)))[0]
    date = (today or datetime.date.today()).strftime("%y%m%d")
    project_dir = os.path.dirname(os.path.abspath(os.fspath(md_path)))
    return Path(project_dir) / "out" / f"{stem}-{date}.docx"

def md2docx(md_path, docx_path, out_path, track_changes=False):
    template_path = os.path.normcase(os.path.abspath(docx_path))
    output_path = os.path.normcase(os.path.abspath(out_path))
    same_file = template_path == output_path
    if not same_file and os.path.exists(out_path):
        same_file = os.path.samefile(docx_path, out_path)
    if same_file:
        raise ValueError("output path must not overwrite the template docx")

    with io.open(md_path, encoding="utf-8") as f:
        md_text = f.read()
    doc = Document(docx_path)
    body_elm = doc.element.body
    anchors = build_anchor_map(doc)
    link_part = doc.part

    blocks = _parse_md_full(md_text)
    anchor_errors = []
    for block in blocks:
        anchor = block["anchor"]
        if not anchor:
            continue
        entry = anchors.get(anchor)
        if entry is None:
            anchor_errors.append(f"{anchor}: not found in template")
            continue
        actual_kind = entry[0]
        expected_kind = "tbl" if block["kind"] == "table" else "p"
        if actual_kind != expected_kind:
            anchor_errors.append(
                f"{anchor}: markdown expects {expected_kind}, template contains {actual_kind}"
            )
    if anchor_errors:
        raise ValueError("invalid markdown anchors: " + "; ".join(anchor_errors))

    cursor = None
    for block in blocks:
        if block["anchor"]:
            entry = anchors.get(block["anchor"])
            kind, elm = entry
            if kind == "p" and block["kind"] != "table":
                p = Paragraph(elm, doc)
                set_paragraph_md(p, block.get("text", ""), link_part,
                                 track=track_changes)
            elif kind == "tbl" and block["kind"] == "table":
                sync_table(Table(elm, doc), block["rows"], link_part,
                           track=track_changes)
            if elm.getparent() is body_elm:  # only body blocks position inserts
                cursor = elm
        else:
            if cursor is None:
                cursor = body_elm[0]
            cursor = insert_block(doc, body_elm, cursor, block, link_part,
                                  track=track_changes)

    os.makedirs(os.path.dirname(os.path.abspath(out_path)) or ".", exist_ok=True)
    doc.save(out_path)
    print(f"rendered: {out_path}")


def main():
    args = [a for a in sys.argv[1:] if a != TRACK_FLAG]
    track = TRACK_FLAG in sys.argv
    if len(args) < 2:
        sys.exit(__doc__)
    md_path, docx_path = args[0], args[1]
    if len(args) > 2:
        out_path = args[2]
    else:
        out_path = default_output_path(md_path, docx_path)
    md2docx(md_path, docx_path, out_path, track_changes=track)


if __name__ == "__main__":
    main()
