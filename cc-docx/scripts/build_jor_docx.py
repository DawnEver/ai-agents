#!/usr/bin/env python3
"""Build Justification of Resources docx — three artifacts:

1. clean     : final document (nested bullets, no marks)
2. review    : INTUITIVE change review — change-log table + legend at top,
               body color-coded by change type (green=new, yellow=edited,
               cyan=moved, none=kept)
3. tracked   : produced separately by compare_docx.py (Word native word-level diff)

Change classification per content item (4th tuple element):
  kept | edited | new | moved
"""
import re, datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = r"C:/Users/linxu/OneDrive - The University of Nottingham/PEMC/260725-proposal-CNT_Winding/260812-JoR/"
DATE = datetime.date.today().strftime("%Y%m%d")

# paragraph shading fills (light tints, readable)
FILL = {"new": "E2EFDA", "edited": "FFF2CC", "moved": "DDEBF7", "kept": None}
CHG_LABEL = {"new": "NEW", "edited": "EDITED", "moved": "MOVED", "kept": "KEPT"}

# --- Content loaded from gitignored workspace module (real names/costs, NOT committed) ---
import os as _os
import sys as _sys
_JOR_DIR = _os.environ.get("JOR_CONTENT_DIR") or _os.path.join(
    _os.path.dirname(_os.path.abspath(__file__)), "..", "workspace", "260812-JoR")
if _JOR_DIR not in _sys.path:
    _sys.path.insert(0, _JOR_DIR)
from jor_content import TITLE, CONTENT, REMOVED

BULLET_STYLE = {"l1": "List Bullet", "l2": "List Bullet 2", "l2p": "List Bullet 2", "l3": "List Bullet 3"}


def shade(p, fill):
    if not fill:
        return
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def new_document():
    doc = Document()
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(11)
    n.paragraph_format.space_after = Pt(4)
    return doc


def add_par(doc, text, style=None, bold=False, centered=False, size=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.font.bold = bold
    if size:
        r.font.size = size
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def build_clean():
    doc = new_document()
    add_par(doc, TITLE, bold=True, centered=True, size=Pt(16))
    for kind, val, *_ in CONTENT:
        if kind == "table":
            rows = val
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = "Table Grid"
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    c = t.cell(i, j)
                    c.text = ""
                    p = c.paragraphs[0]
                    r = p.add_run(cell)
                    r.font.size = Pt(10)
                    if i == 0:
                        r.font.bold = True
        elif kind == "l1":
            add_par(doc, val, style=BULLET_STYLE[kind], bold=True)
        elif kind == "l2":
            add_par(doc, val, style=BULLET_STYLE[kind], bold=True)
        else:
            add_par(doc, val, style=BULLET_STYLE[kind])
    return doc


def build_review():
    """Intuitive review doc: change log + legend on page 1, colour-coded body."""
    doc = new_document()

    # --- Page 1: how to read + change log ---
    add_par(doc, "REVIEW — CHANGES vs ORIGINAL (2026-08-07)", bold=True, size=Pt(14))
    add_par(doc, "Body is colour-coded by change type:")
    for label, fill, meaning in [
        ("NEW", FILL["new"], "content not present in the original"),
        ("EDITED", FILL["edited"], "content rewritten from the original"),
        ("MOVED", FILL["moved"], "content relocated (from MTC/Surrey sections)"),
        ("KEPT", None, "content preserved unchanged"),
    ]:
        p = doc.add_paragraph()
        r = p.add_run(f"■ {label}  = {meaning}")
        if fill:
            shade(p, fill)

    add_par(doc, "", )
    add_par(doc, "Change log:", bold=True)

    log_rows = [["#", "Type", "Item", "What changed"]]
    n = 0
    for kind, val, chg, desc in CONTENT:
        if chg == "kept" or not desc:
            continue
        n += 1
        if kind == "l1":
            label = val
        elif kind == "table":
            label = "Conference cost table"
        else:
            label = (val.split("[")[0].strip()[:44] if "[" in val else val[:44])
        log_rows.append([str(n), CHG_LABEL[chg], label, desc])
    for i, (label, what) in enumerate(REMOVED, start=n + 1):
        log_rows.append([str(i), "REMOVED", label, what])

    t = doc.add_table(rows=len(log_rows), cols=4)
    t.style = "Table Grid"
    for i, row in enumerate(log_rows):
        for j, cell in enumerate(row):
            c = t.cell(i, j)
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(cell)
            r.font.size = Pt(8.5)
            if i == 0 or log_rows[i][1] in ("NEW", "EDITED", "MOVED", "REMOVED"):
                r.font.bold = (i == 0)
            if i > 0:
                fill = {"NEW": FILL["new"], "EDITED": FILL["edited"],
                        "MOVED": FILL["moved"], "REMOVED": "E7E6E6"}.get(log_rows[i][1])
                if fill:
                    shade(p, fill)
    doc.add_page_break()

    # --- Body (colour-coded) ---
    add_par(doc, TITLE, bold=True, centered=True, size=Pt(16))
    for kind, val, chg, _desc in CONTENT:
        if kind == "table":
            rows = val
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = "Table Grid"
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    c = t.cell(i, j)
                    c.text = ""
                    p = c.paragraphs[0]
                    r = p.add_run(cell)
                    r.font.size = Pt(10)
                    if i == 0:
                        r.font.bold = True
                    if chg != "kept":
                        shade(p, FILL[chg])
        elif kind in ("l1", "l2", "l2p", "l3"):
            p = add_par(doc, val, style=BULLET_STYLE[kind], bold=(kind in ("l1", "l2")))
            if chg != "kept":
                shade(p, FILL[chg])
    return doc


def word_count():
    joined = [TITLE] + [val for kind, val, *_ in CONTENT if kind != "table"] \
             + [" ".join(r) for kind, val, *_ in CONTENT if kind == "table" for r in val]
    return len(re.findall(r"\b\w+\b", " ".join(joined)))


if __name__ == "__main__":
    build_clean().save(f"{BASE}JoR _{DATE}_draft.docx")
    build_review().save(f"{BASE}JoR _{DATE}_draft_review.docx")
    print(f"clean:  {BASE}JoR _{DATE}_draft.docx")
    print(f"review: {BASE}JoR _{DATE}_draft_review.docx")
    print(f"word count (JoR body incl. table cells): {word_count()}")
